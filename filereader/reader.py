from docx import Document
from pypdf import PdfReader
from io import BytesIO
from fastapi import UploadFile

import logging
logger = logging.getLogger(__name__)

def extract_text_from_pdf(file_obj):
    # 判断是否是文件路径
    if isinstance(file_obj, str):
        reader = PdfReader(file_obj)
    else:
        file_obj.seek(0)
        bio = BytesIO(file_obj.read())
        reader = PdfReader(bio)
    text = ""
    try:
        for page in reader.pages:
            text += page.extract_text() + "\n"
    except Exception as e:
        logger.error(f"PDF 文本提取失败: {e}")
        logger.debug("pdf_file:", file_obj)
    return text

def extract_text_from_docx(file_obj):
    # 判断是否是路径
    if isinstance(file_obj, str):
        doc = Document(file_obj)
    else:
        file_obj.seek(0)
        bio = BytesIO(file_obj.read())
        doc = Document(bio)
    try:
        text = "\n".join(p.text for p in doc.paragraphs)
    except Exception as e:
        logger.error(f"DOCX 文本提取失败: {e}")
        logger.debug("docx_file:", file_obj)
        text = ""
    return text

def chunking(text, chunk_size=1024):
    """
    将文本切分成指定大小的块
    """
    chunks = []
    # 主要按字符数切分，在切分时寻找周围最近的标点符号，并将当前chunk切分剩余部分保留到下一个chunk
    # 标点符号列表
    punctuation = "。！？；.;?!"
    last_chunk = ""
    for i in range(0, len(text), chunk_size):
        initial_chunk = last_chunk + text[i:i+chunk_size]
        last_chunk = ""
        # 检查是否包含标点符号
        if any(p in initial_chunk for p in punctuation):
            # 找到最近的标点符号位置
            last_punct = max(initial_chunk.rfind(p) for p in punctuation)
            if last_punct != -1:
                # 切分在标点符号后面
                chunk = initial_chunk[:last_punct+1]
                # 保留未处理的部分到下一个chunk
                last_chunk = initial_chunk[last_punct+1:]
        else:
            # 如果没有标点符号，直接添加到chunk
            chunk = initial_chunk
        chunks.append(chunk)
    return chunks

def get_text_from_input(
        message: str | None,
        file: UploadFile | None = None,
) -> str:
    if message:
        return message
    elif file:
        file_ext = file.filename.split(".")[-1]
        if file_ext == "pdf":
            return extract_text_from_pdf(file.file)
        elif file_ext == "docx":
            return extract_text_from_docx(file.file)
        else:
            logger.error(f"不支持的文件格式: {file_ext}")
            return ""
    else:
        logger.error("未提供消息或文件")
        return ""
